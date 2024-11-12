import os
import requests
from bs4 import BeautifulSoup
import csv
from docx import Document
import string
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from datetime import datetime

# Function to get Google Custom Search results
def get_search_results(query, api_key, search_engine_id):
    search_url = f'https://www.googleapis.com/customsearch/v1?q={query}&key={api_key}&cx={search_engine_id}'
    response = requests.get(search_url)
    
    if response.status_code == 200:
        return response.json().get('items', [])
    else:
        print(f"API Error (Status Code {response.status_code}): {response.text}")
        return []

# Function to fetch article content and filter out non-relevant parts
def fetch_article_content(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Attempt to find the main article body, adjusting for different websites
        article_body = soup.find_all('p')
        content = ' '.join([para.get_text().strip() for para in article_body if para.get_text().strip()])
        
        # Optional: Further filtering could be added here to exclude specific unwanted sections (ads, footers)
        return content
    else:
        print(f"Failed to retrieve the page: {url}")
        return ""

# Function to preprocess text: remove punctuation, convert to lowercase, and remove stop words
def preprocess_text(text):
    text = text.lower()
    text = ''.join(char for char in text if char not in string.punctuation)
    return text

# Function to extract unique, relevant sentences (focusing on facts and insights)
def extract_unique_points(articles):
    all_sentences = []
    for article in articles:
        sentences = article.split('.')
        sentences = [s.strip() for s in sentences if len(s.strip()) > 0]  # Clean up sentences
        all_sentences.extend(sentences)

    # Create a TF-IDF Vectorizer to convert sentences into vectors
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(all_sentences)
    
    # Compute cosine similarities to compare sentences and detect repetition
    cosine_sim = cosine_similarity(tfidf_matrix)
    
    unique_sentences = []
    threshold = 0.2  # Lower threshold to determine uniqueness (you can adjust this)
    
    for idx, sentence in enumerate(all_sentences):
        # Check how similar the current sentence is to the others
        similarity = cosine_sim[idx]
        
        # If this sentence is not too similar to others, consider it unique
        if all(similarity_score < threshold for i, similarity_score in enumerate(similarity) if i != idx):
            unique_sentences.append(sentence)
    
    return unique_sentences

# Function to save the results to a Word document
def save_to_word_document(unique_points, folder_path):
    try:
        # Create the document
        doc = Document()
        doc.add_heading('Unique Points from Articles', 0)
        doc.add_paragraph('The following are the unique points extracted from the articles:\n')
        
        for point in unique_points:
            doc.add_paragraph(point)

        # Define the file name and path
        base_file_name = 'article_unique_points.docx'
        file_path = os.path.join(folder_path, base_file_name)
        
        # If the file exists, append a timestamp or number to make the file name unique
        if os.path.exists(file_path):
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_path = os.path.join(folder_path, f'article_unique_points_{timestamp}.docx')

        # Ensure folder exists
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        # Save the document
        doc.save(file_path)
        print(f"Document saved to {file_path}")
    except Exception as e:
        print(f"Error while saving document: {e}")

# Google API setup (Remember to keep these secret!)
API_KEY = 'AIzaSyB9JelqrNGjbW087iNF6Ps7Oa6IINuri2c'
SEARCH_ENGINE_ID = 'a165e30b792d04d08'

url = 'https://www.thehindu.com/news/international/netanyahu-says-he-okayed-lebanon-pager-attacks/article68853655.ece'
custom_folder_path = r"D:\Desktop\New folder"

# Scraping the article
try:
    response = requests.get(url)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, 'html.parser')
    title = soup.title.get_text().strip()
    paragraphs = soup.find_all('p')

    summary = []
    for paragraph in paragraphs:
        text = paragraph.get_text().strip()
        if text:
            summary.append(text)
        if len(summary) >= 3:
            break

    with open('article_summary.csv', mode='w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(['Title', 'Summary'])
        writer.writerow([title, ' '.join(summary)])

    print("Data saved to article_summary.csv")

    # Fetch search results based on the article title
    search_results = get_search_results(title, API_KEY, SEARCH_ENGINE_ID)

    if search_results:
        articles = []
        for item in search_results:
            article_url = item['link']
            print(f"Fetching content from: {article_url}")
            article_content = fetch_article_content(article_url)
            if article_content:
                articles.append(article_content)

        # Extract unique points from the articles
        unique_points = extract_unique_points(articles)
        save_to_word_document(unique_points, custom_folder_path)
    else:
        print("No similar articles found.")

except requests.exceptions.RequestException as e:
    print(f"An error occurred: {e}")

